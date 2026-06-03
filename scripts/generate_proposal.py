#!/usr/bin/env python3
"""Generate formal 开题报告 .docx based on research topic and outline."""

import datetime
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_run_font(run, font_name_cn, font_name_en, size_pt, bold=False, color=None):
    """Set font properties for a run."""
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.font.name = font_name_en
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_heading_styled(doc, text, level=1):
    """Add a heading with Chinese academic formatting."""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = 'Times New Roman'
        r = run._element
        r.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        if level == 0:
            run.font.size = Pt(22)
        elif level == 1:
            run.font.size = Pt(16)
        elif level == 2:
            run.font.size = Pt(14)
        elif level == 3:
            run.font.size = Pt(13)
    return heading

def add_body_para(doc, text, indent=True, bold=False, size=12):
    """Add body text paragraph with Chinese academic formatting."""
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    para.paragraph_format.space_after = Pt(6)
    if indent:
        para.paragraph_format.first_line_indent = Cm(0.74)  # ~2 Chinese chars
    run = para.add_run(text)
    set_run_font(run, '宋体', 'Times New Roman', size, bold=bold)
    return para

def add_table_row(table, cells_text, header=False, widths=None):
    """Add a row to table with formatting."""
    row = table.add_row()
    for i, text in enumerate(cells_text):
        cell = row.cells[i]
        cell.text = ''
        para = cell.paragraphs[0]
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = para.add_run(text)
        if header:
            set_run_font(run, '黑体', 'Times New Roman', 11, bold=True)
            set_cell_shading(cell, 'D9E2F3')
        else:
            set_run_font(run, '宋体', 'Times New Roman', 10.5)
    return row

def main():
    doc = Document()

    # --- Page setup ---
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # --- Style tweaks ---
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    # ============================================================
    # COVER PAGE
    # ============================================================
    for _ in range(6):
        doc.add_paragraph()

    # Title
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run('硕 士 学 位 论 文 开 题 报 告')
    set_run_font(run, '黑体', 'Times New Roman', 26, bold=True)

    doc.add_paragraph()
    doc.add_paragraph()

    # Subtitle
    sub_para = doc.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run('人工智能引发的职业认同威胁对知识型员工\n工作投入的影响：心理资本的中介作用与组织支持感的调节作用')
    set_run_font(run, '黑体', 'Times New Roman', 18, bold=True)

    doc.add_paragraph()
    doc.add_paragraph()

    # Info table
    info_table = doc.add_table(rows=0, cols=2)
    info_table.style = 'Table Grid'
    info_data = [
        ('学科专业', '应用心理学'),
        ('研究方向', '组织心理学 / 职业健康心理学'),
        ('学位类型', '学术型硕士'),
        ('指导教师', '（待填写）'),
        ('学生姓名', '（待填写）'),
        ('学    号', '（待填写）'),
        ('填表日期', datetime.date.today().strftime('%Y 年 %m 月 %d 日')),
    ]
    for label, value in info_data:
        row = info_table.add_row()
        # Label cell
        cell0 = row.cells[0]
        cell0.text = ''
        p0 = cell0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(label)
        set_run_font(r0, '黑体', 'Times New Roman', 14, bold=True)
        cell0.width = Cm(4)
        # Value cell
        cell1 = row.cells[1]
        cell1.text = ''
        p1 = cell1.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(value)
        set_run_font(r1, '宋体', 'Times New Roman', 14)
        cell1.width = Cm(10)

    doc.add_page_break()

    # ============================================================
    # TABLE OF CONTENTS (manual, brief)
    # ============================================================
    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = toc_title.add_run('目  录')
    set_run_font(run, '黑体', 'Times New Roman', 16, bold=True)

    doc.add_paragraph()

    toc_items = [
        '一、研究背景与意义 ................................................. 3',
        '二、国内外研究现状 ................................................. 5',
        '三、研究内容与目标 ................................................. 9',
        '四、研究方法与技术路线 ............................................ 10',
        '五、研究的创新点 .................................................. 12',
        '六、预期成果 ...................................................... 13',
        '七、研究进度安排 .................................................. 14',
        '八、参考文献 ...................................................... 15',
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        set_run_font(run, '宋体', 'Times New Roman', 12)

    doc.add_page_break()

    # ============================================================
    # 一、研究背景与意义
    # ============================================================
    add_heading_styled(doc, '一、研究背景与意义', level=1)

    add_heading_styled(doc, '1.1 研究背景', level=2)

    add_body_para(doc,
        '人工智能（AI）正以前所未有的速度和广度渗透进入职场。从 ChatGPT 到 Midjourney，'
        '从 GitHub Copilot 到智能客服系统，AI 工具已从"辅助工具"演变为"认知劳动的参与者"。'
        'Goldman Sachs（2023）发布的报告估计，全球约 3 亿个全职岗位可能受到 AI 自动化的影响。'
        '普华永道（PwC, 2022）全球劳动力调查显示，约三分之一的受访员工担心 AI 将在短期内替代自己的工作。')

    add_body_para(doc,
        '与历次技术革命不同——前三次工业革命主要替代的是体力劳动——这一轮 AI 浪潮直指知识型劳动者'
        '的核心领域：写作、编程、设计、分析、决策。当 AI 开始完成那些曾被视为"专业壁垒"的认知任务时，'
        '知识型员工面临的不仅是"工作是否会被替代"的生存焦虑，更是一种更深层的心理冲击——'
        '"我多年来积累的专业技能还有价值吗？"这种冲击直指职业认同（professional identity）的核心。')

    add_body_para(doc,
        '中国人民大学心理学系 2024 年对全国 4,420 名产业工人的大规模调查发现，25-30 岁、高学历群体'
        '对 AI 的威胁感知最为强烈，且存在"既认可 AI 效用，又害怕被替代"的矛盾心态（辛自强 & 董妍, 2025）。'
        '这提示我们，AI 对职场的心理冲击并非均匀分布，知识水平越高的群体可能承受越强的认同威胁。')

    add_body_para(doc,
        '现有研究对 AI 职场心理影响的关注集中在"焦虑"和"不安全感"等情绪层面，'
        '但较少追问这些情绪背后的认知机制——即 AI 如何动摇了员工的职业身份感。'
        'Petriglieri（2011）在 Academy of Management Review 上提出的身份威胁理论'
        '（Identity Threat Theory）为理解这一问题提供了有力的分析框架：'
        '当个体感知到某一身份的价值、意义或表现可能受到损害时，会产生一系列防御性和适应性的心理反应。'
        '然而，这一理论尚未被系统应用于 AI 冲击下的职场心理研究。')

    add_heading_styled(doc, '1.2 研究意义', level=2)

    add_heading_styled(doc, '1.2.1 理论意义', level=3)

    add_body_para(doc,
        '第一，将 Petriglieri（2011）身份威胁理论引入 AI 组织行为研究领域，拓展该理论在技术变革情境中的适用边界。'
        '第二，首次在 AI 语境下应用 George et al.（2023）发表于 Journal of Applied Psychology 的身份威胁量表，'
        '检验其在中国知识型员工群体中的适用性，并探索 AI 特有的认同威胁维度。'
        '第三，构建"认同威胁 → 心理资本消耗 → 工作投入下降"的完整心理机制链，'
        '同时考察组织支持感的缓冲效应，丰富资源保存理论（COR）在 AI 情境下的应用。')

    add_heading_styled(doc, '1.2.2 实践意义', level=3)

    add_body_para(doc,
        '第一，为企业 AI 导入过程中的员工心理支持方案提供循证依据，帮助企业识别和保护员工的职业认同资源。'
        '第二，为管理者提供具体的"组织支持行为清单"——什么样的管理举措能够有效缓冲 AI 带来的身份威胁感知。'
        '第三，为员工自身的职业发展和心理调适提供参考路径，特别是心理资本的可开发性为干预提供了理论基础。')

    doc.add_page_break()

    # ============================================================
    # 二、国内外研究现状
    # ============================================================
    add_heading_styled(doc, '二、国内外研究现状', level=1)

    add_heading_styled(doc, '2.1 人工智能对职场的心理影响', level=2)

    add_body_para(doc,
        'AI 对员工心理的影响已成为组织行为学领域的热点议题。传统研究主要沿"技术压力"'
        '（technostress）框架展开，Tarafdar et al.（2007）提出了五维度技术压力模型，'
        '涵盖技术过载、技术入侵、技术复杂性、技术不安全感和技术不确定性。'
        '然而，Sapkota et al.（2025）的最新系统综述指出，传统技术压力框架已不足以捕捉 AI 特有的心理效应，'
        '并提出了五种"AI-stressors"：技术不可预测性、自主性丧失、伦理冲突、社会侵蚀和职业断裂。')

    add_body_para(doc,
        'Rizkina et al.（2025）的 PRISMA 系统综述将 AI 引发的职场焦虑归纳为五种相互关联的形式：'
        '技术焦虑、AI 焦虑、工作替代焦虑、工作不安全感和未来焦虑。'
        '在实证研究方面，Tang et al.（2023）发表于 Journal of Applied Psychology 的研究发现，'
        '与 AI 而非人类同事互动会引发员工的双刃剑效应——既增加对他人的助人行为（适应性），'
        '也引发孤独感、失眠和饮酒增加（非适应性），且对高依恋焦虑的个体影响更强。')

    add_body_para(doc,
        '在中国情境下，Hou & Fan（2024）以五星级酒店员工为对象，基于资源保存理论（COR）'
        '和社会支持理论，首次在中国验证了"AI 压力 → 心理资本消耗 → 工作投入下降"的路径，'
        '并发现组织支持感能够缓冲这一负面过程。该研究发表在 Behavioral Sciences 期刊，'
        '为 AI 职场心理的本土化研究提供了重要的实证基础，但其研究对象限于酒店行业，'
        '未能涵盖知识型白领群体。Wu, Liang & Wang（2024）则在 Journal of Business and Psychology 上'
        '发表研究，发现正念（mindfulness）能够缓冲人机协作不安全感对创造性绩效和幸福感的负面影响。')

    add_heading_styled(doc, '2.2 职业认同威胁', level=2)

    add_body_para(doc,
        '身份威胁理论由 Petriglieri（2011）在 Academy of Management Review 上系统提出。'
        '她将个体层面的身份威胁定义为"被评估为可能损害某一身份的价值、意义或表现的体验"，'
        '并将威胁分为三个相互关联的维度：（1）对身份价值的威胁——感到该身份在他人眼中被贬低或失去价值；'
        '（2）对身份意义的威胁——对该身份"意味着什么"感到不确定或动摇；'
        '（3）对身份表现的威胁——感知到无法表达、执行或活出该身份的能力受到限制。')

    add_body_para(doc,
        'George et al.（2023）在 Journal of Applied Psychology 上发表了基于上述理论框架的身份威胁量表，'
        '经过多轮严格的心理测量学检验（EFA + CFA），最终形成 19 个题项、三个维度的标准化工具。'
        '该量表的贡献在于将身份威胁从一个理论概念转化为可操作的测量工具，'
        '但其开发验证情境主要是教师面对教育技术变革和孕期领导者面对职场偏见——'
        '尚未被应用于 AI 技术变革这一新兴情境。')

    add_body_para(doc,
        '在 AI 与身份威胁的交叉领域，Ziegelmayer & James（2024）提出使用生成式 AI 工具'
        '可能引发职业内疚感——当 AI 完成那些构成职业身份核心的任务（如写作、编程、设计）时，'
        '员工会质疑"我还是一个真正的[写作者/程序员/设计师]吗？"'
        'Hicks & Hevesi（2024）在软件开发者的研究中提出了"AI 技能威胁"概念，'
        '发现种族少数群体、女性和 LGBTQ+开发者感受到更高水平的 AI 技能威胁。')

    add_heading_styled(doc, '2.3 心理资本', level=2)

    add_body_para(doc,
        '心理资本（Psychological Capital, PsyCap）是 Luthans et al.（2007）基于积极组织行为学提出的核心构念，'
        '定义为"个体积极的心理发展状态"，包含四个维度：自我效能（自信）、希望、乐观和韧性。'
        'PCQ-24 量表已被广泛验证，其中文版（李超平译）在中国组织心理学研究中有大量应用。'
        '心理资本的中介角色已在多种压力情境下得到验证——在压力源与员工结果之间，'
        '心理资本充当着重要的心理资源传导机制。'
        'Hou & Fan（2024）首次将 PsyCap 应用于 AI 职场压力研究，发现 AI 压力通过消耗心理资本间接降低工作投入。')

    add_heading_styled(doc, '2.4 组织支持感', level=2)

    add_body_para(doc,
        '组织支持感（POS）是 Eisenberger et al.（1986）基于社会交换理论提出的构念，'
        '指员工对组织重视其贡献、关心其福祉的总体感知。'
        '大量研究证实 POS 是压力源与负面结果之间的重要缓冲变量。'
        '在 AI 情境下，Hou & Fan（2024）发现 POS 显著调节了 AI 压力与心理资本的关系——'
        '当员工感知到高水平的组织支持时，AI 压力对心理资本的消耗效应显著减弱。')

    add_heading_styled(doc, '2.5 研究缺口与本研究切入点', level=2)

    add_body_para(doc,
        '综合以上文献，识别出以下四个清晰的研究缺口：', bold=True)

    add_body_para(doc,
        '缺口一（理论视角）：现有 AI 职场心理研究以"焦虑"和"不安全感"的现象层为主，'
        '缺乏"身份威胁"的机制层视角。Petriglieri（2011）的身份威胁理论为重新理解'
        'AI 的心理影响提供了尚未被开发的理论杠杆。')

    add_body_para(doc,
        '缺口二（测量工具）：George et al.（2023）的 JAP 身份威胁量表尚未在 AI 语境下被应用或修订。'
        'AI 作为一种"智能体"而非"机器"的独特属性，可能引发传统身份威胁框架未能覆盖的心理现象。')

    add_body_para(doc,
        '缺口三（研究对象）：中国现有 AI 职场心理实证研究集中在酒店员工（Hou & Fan, 2024）'
        '和产业工人（辛自强 & 董妍, 2025），知识型白领——这个承受 AI 冲击最直接、最强烈的群体——'
        '几乎是一个研究空白。')

    add_body_para(doc,
        '缺口四（研究方法）：现有研究几乎全部采用横截面问卷设计，受到共同方法偏差和因果推断的局限。'
        '混合方法设计（质性探索 + 定量验证）在 AI 职场心理研究中极为罕见。')

    add_body_para(doc,
        '本研究正是在上述四个缺口的交汇处定位自己的独特贡献。')

    doc.add_page_break()

    # ============================================================
    # 三、研究内容与目标
    # ============================================================
    add_heading_styled(doc, '三、研究内容与目标', level=1)

    add_heading_styled(doc, '3.1 研究问题', level=2)

    add_body_para(doc, '本研究围绕三个核心研究问题展开：')

    add_body_para(doc,
        'Q1（探索性）：在中国知识型员工的感知中，AI 引发了哪些具体的职业认同威胁？'
        '这些威胁的内容、触发情境和心理体验是什么？', bold=True)

    add_body_para(doc,
        'Q2（验证性）：AI 引发的职业认同威胁是否通过心理资本的中介作用影响工作投入和创造力？'
        '这一间接效应的强度如何？', bold=True)

    add_body_para(doc,
        'Q3（边界条件）：组织支持感是否能够缓冲 AI 职业认同威胁对心理资本的消耗效应？'
        '即是否存在有调节的中介效应？', bold=True)

    add_heading_styled(doc, '3.2 研究模型', level=2)

    add_body_para(doc,
        '本研究构建了一个有调节的中介模型（moderated mediation model）。'
        '自变量为 AI 引发的职业认同威胁（三维度：价值威胁、意义威胁、表现威胁），'
        '因变量为工作投入和创造力，中介变量为心理资本（四维度），'
        '调节变量为组织支持感。模型同时检验：(1) 心理资本的中介效应（H2），'
        '以及 (2) 组织支持感对该中介效应的调节作用（H5）。')

    add_heading_styled(doc, '3.3 研究假设', level=2)

    hypotheses = [
        ('H1a', 'AI 引发的职业认同威胁显著负向预测工作投入'),
        ('H1b', 'AI 引发的职业认同威胁显著负向预测创造力'),
        ('H2a', '心理资本在 AI 职业认同威胁与工作投入之间起中介作用'),
        ('H2b', '心理资本在 AI 职业认同威胁与创造力之间起中介作用'),
        ('H3', 'AI 职业认同威胁显著负向预测心理资本的四个维度（自我效能/希望/乐观/韧性）'),
        ('H4', '组织支持感调节 AI 职业认同威胁与心理资本的关系——组织支持感越高，负向影响越弱'),
        ('H5', '组织支持感调节心理资本在 AI 职业认同威胁与结果变量之间的中介效应（有调节的中介）'),
    ]

    for code, hypothesis in hypotheses:
        add_body_para(doc, f'{code}：{hypothesis}。')

    add_heading_styled(doc, '3.4 研究目标', level=2)

    add_body_para(doc,
        '目标一：通过半结构化访谈，探索中国知识型员工感知到的 AI 引发的职业认同威胁的具体内容、'
        '触发情境和心理体验，为 George et al.（2023）量表的 AI 情境修订提供依据。')

    add_body_para(doc,
        '目标二：通过三波时间滞后问卷，量化检验 AI 引发的职业认同威胁通过心理资本影响'
        '工作投入和创造力的中介路径，以及组织支持感的调节效应。')

    add_body_para(doc,
        '目标三：整合质性与量化发现，提出面向企业实践的心理支持策略建议，'
        '为 AI 时代的知识型员工管理提供循证依据。')

    doc.add_page_break()

    # ============================================================
    # 四、研究方法与技术路线
    # ============================================================
    add_heading_styled(doc, '四、研究方法与技术路线', level=1)

    add_heading_styled(doc, '4.1 研究设计总览', level=2)

    add_body_para(doc,
        '本研究采用顺序性解释设计（Sequential Explanatory Design）——先质后量的混合方法。'
        '研究一（质性）的发现为研究二（定量）的工具修订和假设细化提供依据；'
        '研究二（定量）的结果与研究一（质性）进行三角互证，增强结论的可靠性和丰富性。')

    add_heading_styled(doc, '4.2 研究一：质性探索', level=2)

    add_body_para(doc, '研究一采用半结构化访谈法，具体设计如下：')

    add_body_para(doc,
        '受访者标准：知识型行业从业者（IT/金融/咨询/设计/法律/传媒）；'
        '工作中使用 AI 工具至少 3 个月；年龄 22-45 岁。预计访谈 15-20 人，以理论饱和为停止准则。')

    add_body_para(doc,
        '访谈提纲围绕五个模块设计：(1) 工作日常与 AI 使用情况；(2) AI 使用后的自我感受变化'
        '（自我价值、职业意义、能力自信）；(3) 对同行/同事的观察（投射式提问，降低社会赞许性偏差）；'
        '(4) 面对 AI 冲击的心理调适策略；(5) 公司/领导在 AI 导入过程中的支持行为。')

    add_body_para(doc,
        '数据分析采用 Braun & Clarke（2006）的六步主题分析法，使用 NVivo 辅助编码。'
        '信效度保证策略包括：研究者反思日志、双人编码一致性检验（目标 κ ≥ 0.80）、'
        '参与者反馈验证（member checking）和厚描述。')

    add_heading_styled(doc, '4.3 研究二：量化验证', level=2)

    add_body_para(doc, '研究二采用三波时间滞后问卷设计（间隔 2 周），以降低共同方法偏差并增强因果推断力度。')

    add_body_para(doc,
        '样本计划：T1 发放 380 份（预期 T3 有效样本 N ≥ 300，考虑 20% 流失率）。'
        '招募渠道包括企业 HR 合作、Credamo 见数平台和行业社群。')

    add_body_para(doc,
        '测量工具：(1) AI 引发的职业认同威胁——George et al.（2023）19 题量表，经翻译-回译程序修订；'
        '(2) 心理资本——PCQ-24 中文版（李超平译）；(3) 组织支持感——Eisenberger 8 题简化版；'
        '(4) 工作投入——UWES-9 中文版；(5) 创造力——Zhou & George（2001）13 题量表；'
        '(6) 控制变量——社会赞许性（MCSDS 简版）、AI 使用频率、人口统计学变量。')

    add_body_para(doc,
        '数据分析策略：使用 SPSS 27.0 和 Mplus 8.3。包括 Harman 单因子检验 + ULMC 法评估共同方法偏差；'
        '验证性因子分析检验量表结构效度；Hayes PROCESS 宏（Model 4、Model 1、Model 7）检验中介、'
        '调节和有调节的中介效应，Bootstrap 法（5,000 次重复抽样）估计间接效应的置信区间。')

    add_heading_styled(doc, '4.4 技术路线', level=2)

    # Technical route as a table
    route_table = doc.add_table(rows=0, cols=3)
    route_table.style = 'Table Grid'

    # Header
    header_row = route_table.add_row()
    for i, text in enumerate(['阶段', '任务', '产出']):
        cell = header_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, '黑体', 'Times New Roman', 11, bold=True)
        set_cell_shading(cell, 'D9E2F3')

    # Set column widths
    for row in route_table.rows:
        row.cells[0].width = Cm(3)
        row.cells[1].width = Cm(7)
        row.cells[2].width = Cm(4)

    route_data = [
        ('第 1 月', '文献系统检索与综述\n量表翻译-回译\n访谈提纲设计\n伦理审查', '文献综述初稿\n量表中文版初稿\n访谈提纲定稿'),
        ('第 2 月', '研究一执行\n15-20 人半结构化访谈\n转录与主题分析\n研究模型修正', '访谈转录稿\n主题分析报告\n修正版研究模型'),
        ('第 3 月', '研究二准备与 T1\n问卷设计与预测试（N≈50）\n正式施测 T1（IDT + 控制变量）\nT2 施测（PsyCap + POS）', '预测试报告\nT1 数据\nT2 数据'),
        ('第 4 月', '研究二 T3 与数据分析\nT3 施测（工作投入 + 创造力）\n数据清理与 SEM 分析\n假设检验', '完整数据集\n分析结果\n假设检验结论'),
        ('第 5 月', '论文撰写与修改\n全文写作\n格式排版\n导师审阅 → 修改\n提交', '论文初稿\n论文终稿'),
    ]

    for phase, task, output in route_data:
        row = route_table.add_row()
        for i, text in enumerate([phase, task, output]):
            cell = row.cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(text)
            set_run_font(r, '宋体', 'Times New Roman', 10)

    doc.add_page_break()

    # ============================================================
    # 五、研究的创新点
    # ============================================================
    add_heading_styled(doc, '五、研究的创新点', level=1)

    innovations = [
        (
            '视角创新：从"身份威胁"到"AI 职场心理"的跨界嫁接',
            '区别于现有研究以"焦虑"和"不安全感"为主的现象层视角，本研究将 Petriglieri（2011）'
            '身份威胁理论引入 AI 组织行为研究。这不仅是理论框架的简单移植——AI 作为"智能体"'
            '而非传统"机器"的独特属性，为身份威胁理论提供了全新的拓展情境。'
            '研究一（质性探索）可能发现 AI 特有的认同威胁维度，挑战或丰富现有的三维度框架。'
        ),
        (
            '工具创新：George et al.（2023）JAP 量表在 AI 语境下的首次应用与修订',
            'George et al.（2023）发表于 JAP 的身份威胁量表是目前最权威的测量工具，'
            '但其开发情境（教师面对教育技术、孕期领导者面对偏见）与 AI 职场情境存在本质差异。'
            '本研究将通过质性访谈为量表提供 AI 情境的修订依据，并通过 CFA 检验其在中国知识型员工群体中的适用性。'
        ),
        (
            '方法创新：质性 + 定量的混合方法设计',
            '现有 AI 职场心理研究几乎全部采用横截面问卷设计（Hou & Fan, 2024; Wu et al., 2024），'
            '受到共同方法偏差和因果推断的双重局限。本研究采用顺序性解释设计——'
            '先用访谈深入理解"AI 认同威胁"在中国知识型员工中的具体面貌（研究一），'
            '再用三波时间滞后问卷检验因果路径（研究二），最后通过三角互证丰富结论。'
            '混合方法在此领域的应用极为罕见，具有方法论示范价值。'
        ),
        (
            '群体创新：聚焦知识型白领——填补中国研究的关键空白',
            '中国现有 AI 职场心理实证研究集中在酒店行业和制造业（Hou & Fan, 2024; 辛自强 & 董妍, 2025），'
            '知识型白领——这个与 AI 竞争最直接、身份认同最脆弱的群体——几乎是一个研究空白。'
            '本研究的发现将显著提升 AI 职场心理研究对中国核心城市白领群体的覆盖面和解释力。'
        ),
    ]

    for i, (title, desc) in enumerate(innovations, 1):
        add_body_para(doc, f'创新点 {i}：{title}', bold=True)
        add_body_para(doc, desc)

    doc.add_page_break()

    # ============================================================
    # 六、预期成果
    # ============================================================
    add_heading_styled(doc, '六、预期成果', level=1)

    outcomes = [
        '完成一篇 3-4 万字的硕士学位论文，全面呈现研究的理论框架、方法过程和实证发现。',
        '形成 George et al.（2023）身份威胁量表的 AI 情境中文修订版，为该领域后续研究提供本土化测量工具。',
        '产出一套面向企业管理的"AI 导入期员工心理支持策略建议"，包括管理者行动指南和组织支持行为清单。',
        '力争在研究过程中撰写 1-2 篇可投稿至中文核心期刊（如《心理学报》《心理科学进展》'
        '或《中国人力资源开发》）的学术论文，实现学位论文与期刊发表的双向转化。',
        '建立包含 15-20 份质性访谈转录稿和 300+ 份三波问卷数据的原始资料库，为后续纵向追踪和跨文化比较研究奠定基础。',
    ]

    for i, outcome in enumerate(outcomes, 1):
        add_body_para(doc, f'（{i}）{outcome}')

    doc.add_page_break()

    # ============================================================
    # 七、研究进度安排
    # ============================================================
    add_heading_styled(doc, '七、研究进度安排', level=1)

    schedule = [
        ('第 1 月\n（2025年X月）', '文献系统检索与综述撰写\nGeorge et al. (2023) 量表翻译-回译\n访谈提纲设计\n提交伦理审查申请'),
        ('第 2 月\n（2025年X月）', '研究一：质性研究执行\n招募受访者 15-20 人\n完成半结构化访谈\n转录与主题分析\n修正研究模型与假设'),
        ('第 3 月\n（2025年X月）', '研究二：量化研究（T1-T2）\n问卷设计与预测试（N≈50）\nT1 正式施测（IDT + 人口统计学）\nT2 施测（PsyCap + POS）\n数据初步整理'),
        ('第 4 月\n（2025年X月）', '研究二：量化研究（T3 + 分析）\nT3 施测（工作投入 + 创造力）\n数据合并与清理\n信效度检验（CFA）\n假设检验（SEM + Bootstrap）\n补充分析（多组分析等）'),
        ('第 5 月\n（2025年X月）', '论文撰写与修改\n全文写作（第 1-7 章）\n格式排版与参考文献整理\n提交导师审阅\n根据反馈修改\n论文终稿提交'),
    ]

    sched_table = doc.add_table(rows=0, cols=3)
    sched_table.style = 'Table Grid'

    header_row = sched_table.add_row()
    for i, text in enumerate(['时间节点', '主要任务', '里程碑']):
        cell = header_row.cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        set_run_font(r, '黑体', 'Times New Roman', 11, bold=True)
        set_cell_shading(cell, 'D9E2F3')

    milestones = [
        '文献综述完稿 + 量表修订版定稿',
        '访谈主题分析报告完成',
        'T1 + T2 数据收集完成',
        '数据分析完成 + 假设检验结论',
        '论文终稿提交',
    ]

    for i, (time_str, tasks) in enumerate(schedule):
        row = sched_table.add_row()
        for j, text in enumerate([time_str, tasks, milestones[i]]):
            cell = row.cells[j]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(text)
            set_run_font(r, '宋体', 'Times New Roman', 10)

    doc.add_page_break()

    # ============================================================
    # 八、参考文献
    # ============================================================
    add_heading_styled(doc, '八、参考文献', level=1)

    references = [
        '[1] Eisenberger, R., Huntington, R., Hutchison, S., & Sowa, D. (1986). Perceived organizational support. Journal of Applied Psychology, 71(3), 500–507.',
        '[2] George, M. M., Strauss, K., Mell, J. N., & Vough, H. C. (2023). When "who I am" is under threat: Measures of threat to identity value, meanings, and enactment. Journal of Applied Psychology, 108(12), 1952–1978.',
        '[3] Goldman Sachs. (2023). The potentially large effects of artificial intelligence on economic growth. Global Economics Analyst.',
        '[4] Hobfoll, S. E. (1989). Conservation of resources: A new attempt at conceptualizing stress. American Psychologist, 44(3), 513–524.',
        '[5] Hou, J., & Fan, Z. (2024). Working with AI: The effect of job stress on hotel employees\' work engagement. Behavioral Sciences, 14(11), 1076.',
        '[6] Luthans, F., Youssef, C. M., & Avolio, B. J. (2007). Psychological capital: Developing the human competitive edge. Oxford University Press.',
        '[7] Petriglieri, J. L. (2011). Under threat: Responses to and the consequences of threats to individuals\' identities. Academy of Management Review, 36(4), 641–662.',
        '[8] PwC. (2022). PwC\'s global workforce hopes and fears survey 2022.',
        '[9] Rizkina, H. et al. (2025). Job-related anxiety in the age of artificial intelligence: A systematic review of workplace dynamics. Formosa Journal of Multidisciplinary Research.',
        '[10] Sapkota, P. et al. (2025). From technostressors to AI-stressors: A systematic literature review of stressors associated with AI systems. CEUR Workshop Proceedings.',
        '[11] Schaufeli, W. B., Salanova, M., González-Romá, V., & Bakker, A. B. (2002). The measurement of engagement and burnout: A two sample confirmatory factor analytic approach. Journal of Happiness Studies, 3(1), 71–92.',
        '[12] Tang, P. M., Koopman, J., Mai, K. M., De Cremer, D., et al. (2023). No person is an island: Unpacking the work and after-work consequences of interacting with artificial intelligence. Journal of Applied Psychology, 108(11), 1766–1789.',
        '[13] Tarafdar, M., Tu, Q., Ragu-Nathan, B. S., & Ragu-Nathan, T. S. (2007). The impact of technostress on role stress and productivity. Journal of Management Information Systems, 24(1), 301–328.',
        '[14] Wu, J., Liang, J., & Wang, X. (2024). The buffering role of workplace mindfulness: How job insecurity of human-AI collaboration impacts employees\' work-life-related outcomes. Journal of Business and Psychology, 39(6), 1395–1411.',
        '[15] Zhou, J., & George, J. M. (2001). When job dissatisfaction leads to creativity: Encouraging the expression of voice. Academy of Management Journal, 44(4), 682–696.',
        '[16] Ziegelmayer, J., & James, K. (2024). When GenAI threatens professional identity: A TREOS paper. ECIS 2024 TREOS.',
        '[17] 辛自强, 董妍. (2025). 新时代产业工人的社会心理特点与心理建设. 人民论坛.',
        '[18] 李超平. (译). (2008). 心理资本——打造人的竞争优势. 中国轻工业出版社. (Luthans, F., Youssef, C. M., & Avolio, B. J., 2007)',
    ]

    for ref in references:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.25
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(ref)
        set_run_font(run, '宋体', 'Times New Roman', 10.5)

    # ============================================================
    # SAVE
    # ============================================================
    desktop = os.path.expanduser('~/Desktop')
    output_path = os.path.join(desktop, '开题报告.docx')
    doc.save(output_path)
    print(f'✅ 开题报告已生成：{output_path}')
    return output_path

if __name__ == '__main__':
    main()
