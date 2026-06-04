#!/usr/bin/env python3
"""
生成优化版论文大纲 .docx（含修改注释和附录修改对照表）。

功能：
  1. 读取原始大纲 Markdown + 修改记录 JSON → 生成格式规范的 .docx
  2. 每处修改用灰色楷体注释标注：优先级、原版内容、改后内容、修改理由
  3. 附录：修改对照表（逐章列出原版 vs 优化版差异，含优先级彩色标记）
  4. 无修改记录时仅格式化大纲，输出干净 .docx

依赖：
  - python-docx: pip install python-docx

Usage:
  # 完整模式：大纲 + 修改记录 → 优化版 .docx
  python3 scripts/generate_outline_docx.py 大纲关键词.md \
      --changes changes.json --output outline_v2.docx

  # 仅格式化模式：无修改注释，纯大纲输出
  python3 scripts/generate_outline_docx.py 大纲关键词.md --output outline.docx

  # 输出示例 changes.json 模板
  python3 scripts/generate_outline_docx.py --example-changes > changes.json

  # 指定正文字体（默认楷体，可改为宋体/黑体等）
  python3 scripts/generate_outline_docx.py 大纲.md --changes c.json --font 宋体

修改记录 JSON 格式 (--changes):
  {
    "paper_title": "论文标题",
    "version": "v2.0",
    "summary": {
      "total_sections_before": 66,
      "total_sections_after": 90,
      "new_sections": ["5.4 装配工艺", "5.5 实机验证"],
      "p0_changes": 3, "p1_changes": 5,
      "p2_changes": 8, "p3_changes": 4
    },
    "changes": [
      {
        "chapter": "Ch3 源头控制",
        "section": "3.2 叶轮优化设计",
        "type": "modified",
        "original": "优化设计方法",
        "updated": "六方案迭代设计",
        "reason": "从设计计算报告中提取完整迭代路径",
        "priority": "P1",
        "source_doc": "设计计算报告§3.3-3.8"
      }
    ]
  }
"""
import sys, os, json, argparse, re, time
from pathlib import Path

# ── 颜色常量 ─────────────────────────────────────────────────
# fpdf2 与 python-docx 各自使用不同的颜色表示方式：
#   - python-docx: RGBColor 对象
#   - 此处定义 (R, G, B) 元组，由 set_run_font 统一转换

GRAY          = (128, 128, 128)
DARK_BLUE     = (30, 60, 120)
BLACK         = (0, 0, 0)
WHITE         = (255, 255, 255)
BODY_COLOR    = (50, 50, 50)
HEADER_BG_HEX = "1E3C78"   # 表头深蓝底色（Word shading 用 hex）
ALT_ROW_BG_HEX = "F0F5FF"  # 交替行浅蓝底色

PRIORITY_COLORS = {
    "P0": (200, 40, 40),   # 红色 — 必须修改
    "P1": (220, 140, 40),  # 橙色 — 建议修改
    "P2": (180, 150, 50),  # 黄色 — 可优化
    "P3": (100, 160, 100), # 绿色 — 细节打磨
}

PRIORITY_LABELS = {
    "P0": "🔴 P0 必须修改",
    "P1": "🟠 P1 建议修改",
    "P2": "🟡 P2 可优化",
    "P3": "🟢 P3 细节打磨",
}


# ── 工具函数 ─────────────────────────────────────────────────

def rgb(r: int, g: int, b: int):
    """构造 python-docx RGBColor 对象（延迟导入避免耦合）。"""
    from docx.shared import RGBColor
    return RGBColor(r, g, b)


def set_run_font(run, name: str = "", size=None, bold: bool = False,
                 italic: bool = False, color=None):
    """安全设置 Run 对象的字体属性。color 为 (R,G,B) 元组或 None。"""
    if name:
        run.font.name = name
    if size is not None:
        run.font.size = size
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = rgb(*color)


def add_para(doc, text: str = "", font_name: str = "",
             font_size=None, bold: bool = False, color=None,
             alignment=None, space_after=None, space_before=None):
    """添加段落的便捷函数。

    Returns:
        Paragraph 对象（可继续添加 Run）
    """
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    p = doc.add_paragraph()
    if text:
        run = p.add_run(text)
        set_run_font(run, name=font_name, size=font_size, bold=bold, color=color)
    if alignment is not None:
        p.alignment = alignment
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    if space_before is not None:
        p.paragraph_format.space_before = space_before
    return p


def set_cell_shading(cell, hex_color: str):
    """为表格单元格设置背景底色。"""
    from docx.oxml.ns import nsdecls
    from docx.oxml import parse_xml
    shading = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading)


# ── 大纲解析 ─────────────────────────────────────────────────

def parse_outline_md(filepath: str) -> list[dict]:
    """从 Markdown 文件中解析大纲结构。

    识别 # ~ ###### 标题层级，以及 - 列表项。
    层级 ≤2 视为章节，层级 3 视为子节，列表项挂在最近的父级下。

    Returns:
        [{level: int, title: str, sections: [...], items: [...]}]
    """
    if not os.path.exists(filepath):
        return []

    with open(filepath, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    chapters = []
    current_ch: dict | None = None
    current_sec: dict | None = None

    for line in lines:
        line = line.rstrip()
        if not line.strip():
            continue

        # Markdown 标题
        m = re.match(r'^(#{1,6})\s+(.+)$', line)
        if m:
            level = len(m.group(1))
            title = m.group(2).strip()

            if level <= 2:
                current_ch = {
                    'level': level, 'title': title,
                    'sections': [], 'items': [],
                }
                chapters.append(current_ch)
                current_sec = None
            elif level == 3:
                current_sec = {
                    'level': level, 'title': title, 'items': [],
                }
                if current_ch is not None:
                    current_ch['sections'].append(current_sec)
            continue

        # 列表项
        m2 = re.match(r'^[-*]\s+(.+)$', line)
        if m2:
            item_text = m2.group(1).strip()
            if current_sec is not None:
                current_sec['items'].append(item_text)
            elif current_ch is not None:
                current_ch['items'].append(item_text)

    return chapters


# ── .docx 生成 ──────────────────────────────────────────────

def generate_outline_docx(chapters: list[dict], changes_data: dict,
                          output_path: str, font_name: str = "楷体"):
    """生成优化版大纲 .docx。

    Args:
        chapters: parse_outline_md 的解析结果
        changes_data: 修改记录 JSON 解析后的 dict（可为空 {}）
        output_path: 输出 .docx 路径
        font_name: 正文字体名（需系统已安装）
    """
    from docx import Document
    from docx.shared import Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    title = changes_data.get("paper_title", "论文大纲")
    version = changes_data.get("version", "")
    changes = changes_data.get("changes", [])
    summary = changes_data.get("summary", {})

    # ── 封面 ─────────────────────────────────────

    for _ in range(7):
        doc.add_paragraph()

    add_para(doc, title, font_name=font_name, font_size=Pt(22),
             bold=True, color=DARK_BLUE,
             alignment=WD_ALIGN_PARAGRAPH.CENTER)

    if version:
        add_para(doc, f"优化版 {version}", font_name=font_name,
                 font_size=Pt(12), color=GRAY,
                 alignment=WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc, f"生成时间: {time.strftime('%Y-%m-%d %H:%M')}",
             font_name=font_name, font_size=Pt(9), color=GRAY,
             alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(30))

    # ── 变更摘要 ─────────────────────────────────

    if summary:
        doc.add_page_break()
        add_para(doc, "变更摘要", font_name=font_name, font_size=Pt(16),
                 bold=True, color=DARK_BLUE, space_after=Pt(12))

        before = summary.get("total_sections_before", "—")
        after  = summary.get("total_sections_after", "—")
        new_secs = summary.get("new_sections", [])

        items = [
            f"原大纲子节数: {before}",
            f"优化后子节数: {after}",
            f"新增整节: {', '.join(new_secs) if new_secs else '无'}",
            f"P0（必须修改）: {summary.get('p0_changes', 0)} 项",
            f"P1（建议修改）: {summary.get('p1_changes', 0)} 项",
            f"P2（可优化）:   {summary.get('p2_changes', 0)} 项",
            f"P3（细节打磨）: {summary.get('p3_changes', 0)} 项",
        ]
        for item in items:
            add_para(doc, item, font_name=font_name, font_size=Pt(10.5),
                     color=(60, 60, 60))

    # ── 大纲正文 ─────────────────────────────────

    doc.add_page_break()
    add_para(doc, "优化版大纲", font_name=font_name, font_size=Pt(18),
             bold=True, color=DARK_BLUE, space_after=Pt(16))

    # 构建修改索引: {(章节标题, 子节标题): change_obj}
    change_index: dict[tuple[str, str], dict] = {}
    for c in changes:
        key = (c.get("chapter", ""), c.get("section", ""))
        change_index[key] = c

    for ch in chapters:
        # 章节标题
        add_para(doc, ch['title'], font_name=font_name, font_size=Pt(14),
                 bold=True, color=DARK_BLUE, space_after=Pt(6),
                 space_before=Pt(12))

        for sec in ch.get('sections', []):
            # 子节标题
            p = add_para(doc, font_name=font_name, font_size=Pt(11),
                         color=BLACK, space_after=Pt(2))
            run_title = p.add_run(sec['title'])
            set_run_font(run_title, name=font_name, size=Pt(11),
                         bold=True, color=BLACK)

            # 检查是否有修改记录
            ch_key = (ch['title'], sec['title'])
            change = change_index.get(ch_key)

            if change:
                priority = change.get("priority", "P2")
                original  = change.get("original", "")
                updated   = change.get("updated", "")
                reason    = change.get("reason", "")
                source    = change.get("source_doc", "")

                # 灰色楷体注释标注修改
                note_text = (
                    f"  [{PRIORITY_LABELS.get(priority, priority)}] "
                    f"原: {original}  →  改: {updated}"
                )
                if source:
                    note_text += f"  （来源: {source}）"

                p_note = add_para(doc, font_name=font_name, font_size=Pt(8.5),
                                  color=GRAY, space_after=Pt(1))
                run_note = p_note.add_run(note_text)
                set_run_font(run_note, name=font_name, size=Pt(8.5),
                             italic=True, color=GRAY)

                if reason:
                    p_reason = add_para(doc, font_name=font_name,
                                        font_size=Pt(8.5), color=GRAY,
                                        space_after=Pt(3))
                    run_reason = p_reason.add_run(f"  理由: {reason}")
                    set_run_font(run_reason, name=font_name, size=Pt(8.5),
                                 italic=True, color=GRAY)

            # 子节下的列表项
            for item in sec.get('items', []):
                add_para(doc, f"  • {item}", font_name=font_name,
                         font_size=Pt(10), color=BODY_COLOR)

        # 章节级的列表项
        for item in ch.get('items', []):
            add_para(doc, f"  • {item}", font_name=font_name,
                     font_size=Pt(10), color=BODY_COLOR)

    # ── 附录：修改对照表 ─────────────────────────────

    if not changes:
        # 无修改记录，跳过附录
        doc.save(output_path)
        size_kb = os.path.getsize(output_path) // 1024
        print(f"✅ 大纲已生成（无修改记录，跳过附录）: {output_path} "
              f"({size_kb} KB)", flush=True)
        return output_path

    doc.add_page_break()
    add_para(doc, "附录：修改对照表", font_name=font_name, font_size=Pt(16),
             bold=True, color=DARK_BLUE, space_after=Pt(12))

    add_para(doc,
             f"以下逐章列出原版 v1.0 与优化版 {version or 'v2.0'} 的变更清单。",
             font_name=font_name, font_size=Pt(9), color=GRAY,
             space_after=Pt(10))

    # 创建表格
    table = doc.add_table(rows=1, cols=6)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 表头
    headers = ["章节", "优先级", "类型", "原版", "优化版", "修改理由"]
    header_cells = table.rows[0].cells
    for i, (cell, header) in enumerate(zip(header_cells, headers)):
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(header)
        set_run_font(run, name=font_name, size=Pt(8.5), bold=True, color=WHITE)
        set_cell_shading(cell, HEADER_BG_HEX)

    # 数据行
    for row_idx, change in enumerate(changes):
        row = table.add_row()
        cells = row.cells

        chapter = change.get("chapter", "")
        section = change.get("section", "")
        full_loc = f"{chapter}\n→ {section}" if section else chapter

        values = [
            full_loc,
            change.get("priority", "P2"),
            _type_label(change.get("type", "modified")),
            change.get("original", ""),
            change.get("updated", ""),
            change.get("reason", ""),
        ]

        for i, (cell, val) in enumerate(zip(cells, values)):
            cell.text = ""
            p = cell.paragraphs[0]
            align = WD_ALIGN_PARAGRAPH.CENTER if i in (1, 2) else WD_ALIGN_PARAGRAPH.LEFT
            p.alignment = align
            run = p.add_run(str(val)[:200])
            set_run_font(run, name=font_name, size=Pt(8), color=BLACK)

        # 对于 P0 行，在优先级列加粗并上色
        if change.get("priority") == "P0":
            prio_cell = cells[1]
            for p in prio_cell.paragraphs:
                for run in p.runs:
                    set_run_font(run, name=font_name, size=Pt(8),
                                 bold=True, color=PRIORITY_COLORS["P0"])

        # 交替行底色
        if row_idx % 2 == 1:
            for cell in cells:
                set_cell_shading(cell, ALT_ROW_BG_HEX)

    # 设置列宽（mm → EMU，python-docx 用 Cm）
    widths = [Cm(3.0), Cm(1.0), Cm(1.0), Cm(3.5), Cm(3.5), Cm(4.5)]
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = width

    # ── 保存 ─────────────────────────────────────

    doc.save(output_path)
    size_kb = os.path.getsize(output_path) // 1024
    print(f"✅ 优化版大纲已生成: {output_path} ({size_kb} KB)", flush=True)
    print(f"   含 {len(changes)} 条修改注释 + 附录修改对照表", flush=True)
    return output_path


def _type_label(t: str) -> str:
    """将修改类型转为中文标签。"""
    return {
        "modified": "修改",
        "added": "新增",
        "deleted": "删除",
        "restructured": "重构",
    }.get(t, t)


# ── CLI ──────────────────────────────────────────────────────

def get_example_changes() -> dict:
    """返回示例 changes.json 数据结构。"""
    return {
        "paper_title": "论文标题示例",
        "version": "v2.0",
        "summary": {
            "total_sections_before": 66,
            "total_sections_after": 90,
            "new_sections": ["5.4 装配工艺", "5.5 实机验证"],
            "p0_changes": 3,
            "p1_changes": 5,
            "p2_changes": 8,
            "p3_changes": 4,
        },
        "changes": [
            {
                "chapter": "Ch2 机理分析",
                "section": "2.2 数值模拟方法",
                "type": "modified",
                "original": "数值模拟方法",
                "updated": "k-ε 湍流模型 + SIMPLE 算法 + MRF 模型",
                "reason": "从设计计算报告中提取具体 CFD 设置参数",
                "priority": "P1",
                "source_doc": "设计计算报告 §3.4",
            },
            {
                "chapter": "Ch5 试验验证",
                "section": "5.4 装配工艺",
                "type": "added",
                "original": "（原大纲无此节）",
                "updated": "装配工艺：跳动公差要求、专用工装架设计",
                "reason": "总体技术报告中含完整装配工艺章节，原大纲未覆盖",
                "priority": "P0",
                "source_doc": "总体技术报告 §5.3",
            },
            {
                "chapter": "Ch5 试验验证",
                "section": "5.5 实机验证",
                "type": "added",
                "original": "（原大纲无此节）",
                "updated": "实机换装验证：换装周期、隔振效果实测数据",
                "reason": "改进报告中含实机验证数据，填补工程验证闭环",
                "priority": "P0",
                "source_doc": "改进报告 §15",
            },
            {
                "chapter": "Ch1 绪论",
                "section": "1.1 研究背景",
                "type": "modified",
                "original": "行业背景概述",
                "updated": "行业背景 + 具体产品参数 + 适用标准",
                "reason": "从总体技术报告中提取产品参数和引用标准编号",
                "priority": "P2",
                "source_doc": "总体技术报告 §1",
            },
        ],
    }


def main():
    parser = argparse.ArgumentParser(
        description="生成优化版论文大纲 .docx（含修改注释和附录修改对照表）",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        epilog="""
示例:
  # 完整模式
  %(prog)s 大纲关键词.md --changes changes.json -o outline_v2.docx

  # 仅格式化大纲
  %(prog)s 大纲关键词.md -o outline.docx

  # 输出示例 changes.json 模板
  %(prog)s --example-changes > changes.json
        """,
    )
    parser.add_argument(
        "outline_md", nargs="?",
        help="大纲 Markdown 文件路径",
    )
    parser.add_argument("--changes", "-c",
                        help="修改记录 JSON 文件路径")
    parser.add_argument("--output", "-o", default="outline_v2.docx",
                        help="输出 .docx 文件路径（默认: outline_v2.docx）")
    parser.add_argument("--font", default="楷体",
                        help="正文字体名（默认: 楷体，需系统已安装）")
    parser.add_argument("--example-changes", action="store_true",
                        help="输出示例 changes.json 模板到 stdout 并退出")
    args = parser.parse_args()

    # --example-changes 模式
    if args.example_changes:
        print(json.dumps(get_example_changes(), ensure_ascii=False, indent=2))
        return

    if not args.outline_md:
        parser.print_help()
        print("\n❌ 请指定大纲 Markdown 文件路径", flush=True)
        sys.exit(1)

    if not os.path.exists(args.outline_md):
        print(f"❌ 文件不存在: {args.outline_md}", flush=True)
        sys.exit(1)

    # 检查 python-docx
    try:
        import docx
    except ImportError:
        print("❌ 缺少 python-docx，请执行: pip install python-docx", flush=True)
        sys.exit(1)

    # ── 解析大纲 ──

    chapters = parse_outline_md(args.outline_md)
    if not chapters:
        print("⚠️ 未能从大纲文件中解析出章节结构。", flush=True)
        print("   支持格式: # 章标题 / ## 节标题 / ### 子节 / - 列表项", flush=True)
        sys.exit(1)

    print(f"📖 解析到 {len(chapters)} 个章节:", flush=True)
    for ch in chapters:
        sec_count = len(ch.get('sections', []))
        print(f"   {ch['title']} ({sec_count} 子节)", flush=True)

    # ── 加载修改记录 ──

    changes_data: dict = {}
    if args.changes:
        if os.path.exists(args.changes):
            with open(args.changes, 'r', encoding='utf-8') as f:
                changes_data = json.load(f)
            n = len(changes_data.get("changes", []))
            print(f"📝 加载 {n} 条修改记录", flush=True)
        else:
            print(f"⚠️ 修改记录文件不存在: {args.changes}", flush=True)
            print("   将仅格式化大纲（无修改注释）。", flush=True)

    # ── 生成 ──

    generate_outline_docx(chapters, changes_data, args.output, args.font)


if __name__ == "__main__":
    main()
